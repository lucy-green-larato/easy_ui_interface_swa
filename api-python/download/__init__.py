# /api-python/campaign/download/__init__.py
# GET .docx export built from campaign.json + evidence_log.json
# Usage:
#   GET /api/campaign/download?runId=<id>
#
# Requires in /api-python/requirements.txt:
#   python-docx>=0.8.11

import os
import io
import json
from urllib.parse import urlsplit

import azure.functions as func
from azure.storage.blob import BlobServiceClient
from docx import Document


from function_app import app  # shared HTTP FunctionApp


def _blob_service() -> BlobServiceClient:
    """Build BlobServiceClient from account SAS URL in UPLOADS_SAS_URL."""
    sas_url = os.environ["UPLOADS_SAS_URL"].strip()
    parts = urlsplit(sas_url)
    account_base = f"{parts.scheme}://{parts.netloc}"
    sas_token = parts.query.lstrip("?")
    return BlobServiceClient(account_url=account_base, credential=sas_token)


def _results_container_name() -> str:
    return os.environ["CAMPAIGN_RESULTS_CONTAINER"]


def _find_blob_path(cc, run_id: str, filename: str) -> str:
    """
    Search results/campaign/*/*/*/*/<runId>/<filename>.
    We don't know page/date here, so scan under the fixed prefix.
    """
    prefix = "results/campaign/"
    for b in cc.list_blobs(name_starts_with=prefix):
        if b.name.endswith(f"/{run_id}/{filename}"):
            return b.name
    return ""


def _add_para(doc: Document, text: str):
    doc.add_paragraph(text or "")


def _add_evidence_table(doc: Document, evidence: list):
    if not evidence:
        return
    headers = ["ID", "Publisher", "Title", "Date", "URL"]
    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for item in evidence:
        row = table.add_row().cells
        row[0].text = str(item.get("id", ""))
        row[1].text = str(item.get("publisher", ""))
        row[2].text = str(item.get("title", ""))
        row[3].text = str(item.get("date", ""))
        row[4].text = str(item.get("url", ""))


@app.route(route="campaign/download", methods=["GET"])
def download(req: func.HttpRequest) -> func.HttpResponse:
    run_id = req.params.get("runId")
    if not run_id:
        return func.HttpResponse("Missing runId", status_code=400)

    try:
        bsc = _blob_service()
        container = _results_container_name()
        cc = bsc.get_container_client(container)

        # Load campaign.json (required)
        path_campaign = _find_blob_path(cc, run_id, "campaign.json")
        if not path_campaign:
            return func.HttpResponse("campaign.json not found", status_code=404)
        campaign_bytes = cc.get_blob_client(path_campaign).download_blob().readall()
        campaign = json.loads(campaign_bytes)

        # Load evidence_log.json (optional)
        evidence = []
        path_evidence = _find_blob_path(cc, run_id, "evidence_log.json")
        if path_evidence:
            evidence_bytes = cc.get_blob_client(path_evidence).download_blob().readall()
            evidence = json.loads(evidence_bytes)

        # Build .docx according to the fixed contract
        doc = Document()
        doc.core_properties.title = f"Campaign {run_id}"

        doc.add_heading("Evidence-first Campaign Pack", level=0)
        _add_para(doc, f"Run ID: {run_id}")

        # Executive summary
        exec_sum = campaign.get("executive_summary") or ""
        if exec_sum:
            doc.add_heading("Executive summary", level=1)
            _add_para(doc, exec_sum)

        # Landing page
        lp = campaign.get("landing_page") or {}
        if lp:
            doc.add_heading("Landing page", level=1)
            headline = lp.get("headline")
            subheadline = lp.get("subheadline")
            if headline:
                doc.add_heading("Headline", level=2)
                _add_para(doc, headline)
            if subheadline:
                doc.add_heading("Subheadline", level=2)
                _add_para(doc, subheadline)
            sections = lp.get("sections") or []
            if sections:
                doc.add_heading("Sections", level=2)
                for s in sections:
                    title = s.get("title", "")
                    doc.add_heading(title or "Section", level=3)
                    if "content" in s:
                        _add_para(doc, s.get("content") or "")
                    elif "bullets" in s:
                        for b in (s.get("bullets") or []):
                            _add_para(doc, f"• {b}")
            cta = lp.get("cta")
            if cta:
                doc.add_heading("Call to action", level=2)
                _add_para(doc, cta)

        # Emails
        emails = campaign.get("emails") or []
        if emails:
            doc.add_heading("Emails", level=1)
            for i, em in enumerate(emails, 1):
                subj = em.get("subject", "")
                doc.add_heading(f"Email {i}: {subj}", level=2)
                preview = em.get("preview")
                if preview:
                    _add_para(doc, f"Preview: {preview}")
                _add_para(doc, em.get("body", ""))

        # Sales enablement
        se = campaign.get("sales_enablement") or {}
        if se:
            doc.add_heading("Sales enablement", level=1)
            call_script = se.get("call_script")
            one_pager = se.get("one_pager")
            if call_script:
                doc.add_heading("Call script", level=2)
                _add_para(doc, call_script)
            if one_pager:
                doc.add_heading("One-pager", level=2)
                _add_para(doc, one_pager)

        # Evidence log
        if evidence:
            doc.add_heading("Evidence log", level=1)
            _add_evidence_table(doc, evidence)

        # Input proof
        input_proof = campaign.get("input_proof") or {}
        if input_proof:
            doc.add_heading("Input proof", level=1)
            for k, v in input_proof.items():
                _add_para(doc, f"{k}: {v}")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        return func.HttpResponse(
            body=buf.getvalue(),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="campaign-{run_id}.docx"'},
            status_code=200,
        )

    except KeyError as ke:
        return func.HttpResponse(f"Missing environment variable: {ke}", status_code=500)
    except Exception as e:
        # Hide internal details in production if desired
        return func.HttpResponse(str(e), status_code=500)
